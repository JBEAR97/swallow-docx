import os
import re
from datetime import datetime
from collections import defaultdict

import numpy as np
import pandas as pd
import docx
import time


from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

from flask import Flask, send_file, request
from watchdog.observers import Observer
from watchdog.events import FileSystemEventHandler

app = Flask(__name__)

WATCH_FOLDER = os.environ.get(
    "WATCH_FOLDER", 
    os.path.join(os.path.dirname(os.path.abspath(__file__)), "DOCX")
)
print(f"🔍 DEBUG: WATCH_FOLDER = {WATCH_FOLDER}")  # ← AGGIUNGI QUESTA
print(f"📁 Esiste? {os.path.exists(WATCH_FOLDER)}")

class CorrelationMatrixGenerator:
    def __init__(self, watch_folder):
        self.watch_folder = watch_folder

        # Document-level index
        self.documents = {}              # filename -> full_text
        self.doc_chapters = {}           # filename -> list[(title, content)]
        self.chapter_keys = []           # filenames order in TF-IDF matrix

        # Doc TF-IDF
        self.vectorizer = None
        self.tfidf_matrix = None
        self.correlation_matrix = None
        self.filename_to_index = {}

        self.csv_path = os.path.join(os.path.dirname(self.watch_folder), "output", "swallow_matrix.csv")

        # Word connections (global)
        self.connections = {}            # WORD -> list[(DOC_ID, chapter_title)]
        self.doc_id_to_filename = {}     # DOC_ID -> filename
        self.filename_to_doc_id = {}     # filename -> DOC_ID

        # Chapter-level index (for "Analizza capitolo specifico")
        self.chapter_records = []        # list[dict] with DOC_ID, FILENAME, TITLE, LOCATION
        self.chapter_texts = []          # list[str]
        self.chapter_tfidf_matrix = None # TF-IDF for chapters in same vector space

    # -------------------------
    # Loading / parsing DOCX
    # -------------------------
    def load_all_files(self):
        if not os.path.exists(self.watch_folder):
            print("❌ Cartella non trovata:", self.watch_folder)
            return False

        files = [f for f in os.listdir(self.watch_folder) if f.lower().endswith(".docx")]
        if not files:
            print("⚠ Nessun file .docx trovato.")
            return False

        self.documents.clear()
        self.doc_chapters.clear()
        self.doc_id_to_filename.clear()
        self.filename_to_doc_id.clear()

        for file in files:
            path = os.path.join(self.watch_folder, file)
            try:
                doc = docx.Document(path)
                full_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                self.documents[file] = full_text

                chapters = self.extract_chapters_from_headings(doc)
                self.doc_chapters[file] = chapters

                doc_id = os.path.splitext(file)[0].upper()
                self.doc_id_to_filename[doc_id] = file
                self.filename_to_doc_id[file] = doc_id

                print(f"✅ {file}: {len(chapters)} capitoli trovati")
            except Exception as e:
                print(f"❌ Errore {file}: {e}")

        self.build_chapter_index()

        print(f"📄 Caricati {len(files)} file.")
        return True

    def extract_chapters_from_headings(self, doc):
        """
        Estrae capitoli leggendo gli stili Heading nativi di Word.
        Funziona con: Heading 1/2/3, Intestazione 1/2/3, grassetto, numeri romani.
        """
        chapters = []
        current_title = None
        current_content = []

        heading_styles = {
            'heading 1', 'heading 2', 'heading 3',
            'intestazione 1', 'intestazione 2', 'intestazione 3',
            'titolo 1', 'titolo 2', 'titolo 3',
            'title', 'subtitle'
        }

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name.lower() if para.style else ""
            is_heading = any(h in style_name for h in heading_styles)

            is_bold_title = (
                len(para.runs) > 0
                and all(run.bold for run in para.runs if run.text.strip())
                and 8 < len(text) < 150
            )

            is_roman_title = bool(re.match(r'^[IVXLCDM]+\.\s+.{10,}', text))

            if is_heading or is_bold_title or is_roman_title:
                if current_title:
                    chapters.append((current_title, " ".join(current_content).strip()))
                clean_title = re.sub(r'\[(.+?)\]\(\)', r'\1', text).strip()
                current_title = clean_title
                current_content = []
            else:
                if current_title:
                    current_content.append(text)

        if current_title:
            chapters.append((current_title, " ".join(current_content).strip()))

        return chapters

    def _ignore_titles_set(self):
        return {
            "CULTURAL MONDAY", "GEOGRAPHIC WEDNESDAY",
            "SOCIAL FRIDAY", "TOURISM SUNDAY",
            "TABLE OF CONTENTS", "CONTENTS"
        }

    def build_chapter_index(self):
        """
        Costruisce un indice capitoli: serve per correlazioni "logiche" a livello capitolo.
        """
        ignore_titles = self._ignore_titles_set()

        self.chapter_records = []
        self.chapter_texts = []

        for filename, chapters in self.doc_chapters.items():
            doc_id = os.path.splitext(filename)[0].upper()

            for title, content in chapters:
                title_clean = title.upper().strip()
                if title_clean in ignore_titles:
                    continue
                if "resources" in title_clean.lower() or "sources" in title_clean.lower():
                    continue
                if len(title.strip()) < 3:
                    continue

                text = (content or "").strip()
                if not text:
                    continue

                self.chapter_records.append({
                    "DOC_ID": doc_id,
                    "FILENAME": filename,
                    "TITLE": title,
                    "LOCATION": f"{doc_id}: {title}",
                })
                self.chapter_texts.append(text)

    # -------------------------
    # TF-IDF / similarity
    # -------------------------
    def compute_matrix(self):
        corpus = list(self.documents.values())
        self.chapter_keys = list(self.documents.keys())

        # ✅ N-grammi per mantenere i composti (es. "mount fuji", "climate change", ...)
        self.vectorizer = TfidfVectorizer(
            stop_words="english",
            max_features=20000,
            ngram_range=(1, 3)
        )

        self.tfidf_matrix = self.vectorizer.fit_transform(corpus)
        self.correlation_matrix = cosine_similarity(self.tfidf_matrix)
        self.filename_to_index = {fn: i for i, fn in enumerate(self.chapter_keys)}

        # Chapter matrix nello stesso spazio del vectorizer
        if self.chapter_texts:
            self.chapter_tfidf_matrix = self.vectorizer.transform(self.chapter_texts)
        else:
            self.chapter_tfidf_matrix = None

    def get_doc_similarity(self, base_filename, other_filename):
        if self.correlation_matrix is None:
            return 0.0
        i = self.filename_to_index.get(base_filename)
        j = self.filename_to_index.get(other_filename)
        if i is None or j is None:
            return 0.0
        return float(self.correlation_matrix[i, j])

    # -------------------------
    # CSV matrix export
    # -------------------------
    def generate_csv_matrix(self):
        if not self.documents:
            return None, None

        self.compute_matrix()

        df = pd.DataFrame(
            self.correlation_matrix,
            index=self.chapter_keys,
            columns=self.chapter_keys
        )
        df.to_csv(self.csv_path)
        return df, self.csv_path

    # -------------------------
    # Word connections (global)
    # -------------------------
    def generate_word_connections(self):
        """
        WORD -> lista posizioni (DOC_ID, chapter_title)
        """
        word_map = defaultdict(set)
        ignore_titles = self._ignore_titles_set()

        for filename, chapters in self.doc_chapters.items():
            doc_id = os.path.splitext(filename)[0].upper()

            for title, content in chapters:
                title_clean = title.upper().strip()

                if title_clean in ignore_titles:
                    continue
                if "resources" in title_clean.lower() or "sources" in title_clean.lower():
                    continue
                if len(title.strip()) < 3:
                    continue

                lowered = (content or "").lower()
                words = re.findall(r'\b[a-zA-Z]{3,}\b', lowered)
                for w in words:
                    word_map[w.upper()].add((doc_id, title))

        # Ordinamento: prima le parole che appaiono in più DOC_ID
        def score(item):
            word, locs = item
            doc_count = len(set(d for d, _ in locs))
            occ = len(locs)
            return (doc_count * 10) + occ

        sorted_items = sorted(word_map.items(), key=score, reverse=True)

        self.connections = {
            word: list(locs)
            for word, locs in sorted_items
        }

        print(f"🔑 {len(self.connections)} parole indicizzate (globale).")
        return self.connections

    # -------------------------
    # Multi-word semantic query (document-level)
    # -------------------------
    def search_query_connections(self, query, min_sim=0.10, top_n=5):
        """
        Cerca capitoli che contengono tutte le parole (split classico), poi ranking doc-level.
        """
        if not query or self.tfidf_matrix is None or self.correlation_matrix is None:
            return []

        terms = re.findall(r'\b[a-zA-Z]{3,}\b', query.lower())
        if not terms:
            return []

        ignore_titles = self._ignore_titles_set()
        results = []

        for filename, chapters in self.doc_chapters.items():
            doc_id = os.path.splitext(filename)[0].upper()

            for title, content in chapters:
                title_clean = title.upper().strip()
                if title_clean in ignore_titles:
                    continue
                if "resources" in title_clean.lower() or "sources" in title_clean.lower():
                    continue

                lowered = (content or "").lower()
                if not all(t in lowered for t in terms):
                    continue

                idx = self.filename_to_index.get(filename)
                if idx is None:
                    continue

                row = self.correlation_matrix[idx]
                sim_indices = sorted(range(len(row)), key=lambda i: row[i], reverse=True)

                top_related = []
                for i in sim_indices:
                    if i == idx:
                        continue
                    sim = float(row[i])
                    if sim < min_sim:
                        continue
                    top_related.append({"DOCUMENTO": self.chapter_keys[i], "SIMILARITÀ": sim})
                    if len(top_related) >= top_n:
                        break

                results.append({
                    "QUERY": query,
                    "PAROLE_QUERY": ", ".join(terms),
                    "PAESE_CAPITOLO": f"{doc_id}: {title}",
                    "DOCUMENTO_BASE": filename,
                    "CORRELAZIONI": top_related
                })

        return results

    # -------------------------
    # Chapter-specific phrase logic (NO splitting compositions)
    # -------------------------
    def analyze_chapter_specific_phrase_logic(
        self,
        chapter_query: str,
        phrase_query: str = "",
        top_n_links: int = 12,
        min_score: float = 0.06,
        alpha_phrase: float = 0.70
    ):
        """
        - Trova un capitolo base usando chapter_query (substring su "DOC: titolo" o titolo).
        - phrase_query è trattata come FRASE/COMPOSTO (non viene tokenizzata manualmente).
          Puoi passare più frasi separate da ';' o newline: ogni frase produce un blocco output.
        - Ranking "concettuale":
            score = alpha * sim(phrase, chapter) + (1-alpha) * sim(base_chapter, chapter)
        - Output: list[report] con righe LOCATION ordinate per score.
        """
        if not chapter_query or self.chapter_tfidf_matrix is None:
            return None

        cq = chapter_query.lower().strip()

        # 1) trova capitolo base
        base_idx = None
        for i, rec in enumerate(self.chapter_records):
            searchable = rec["LOCATION"].lower()
            if cq in searchable or cq in rec["TITLE"].lower():
                base_idx = i
                break

        if base_idx is None:
            return None

        base_location = self.chapter_records[base_idx]["LOCATION"]
        base_vec = self.chapter_tfidf_matrix[base_idx]

        # 2) frasi (senza scorporare)
        raw = (phrase_query or "").strip()
        if raw:
            phrases = [p.strip() for p in re.split(r"[;\n]+", raw) if p.strip()]
        else:
            phrases = [""]  # se vuota: ranking solo per vicinanza al capitolo base

        # sim(base, chapter) per tutti i capitoli
        sim_base = cosine_similarity(base_vec, self.chapter_tfidf_matrix).flatten()

        reports = []

        for ph in phrases:
            label = ph.upper() if ph else "(BASE CHAPTER RELEVANCE)"

            if ph:
                q_vec = self.vectorizer.transform([ph])  # la frase rimane intera
                sim_phrase = cosine_similarity(q_vec, self.chapter_tfidf_matrix).flatten()
            else:
                sim_phrase = np.zeros_like(sim_base)

            score = (alpha_phrase * sim_phrase) + ((1.0 - alpha_phrase) * sim_base)

            idxs = np.argsort(-score)

            lines = []
            lines.append({"LOCATION": base_location, "SCORE": 1.0})

            added = {base_location}
            for j in idxs:
                loc = self.chapter_records[int(j)]["LOCATION"]
                if loc in added:
                    continue
                s = float(score[int(j)])
                if s < min_score:
                    continue
                lines.append({"LOCATION": loc, "SCORE": s})
                added.add(loc)
                if len(lines) >= top_n_links:
                    break

            reports.append({
                "LABEL": label,
                "BASE": base_location,
                "LINES": lines
            })

        return reports


class DocxChangeHandler(FileSystemEventHandler):
    def __init__(self, generator):
        self.generator = generator

    def on_any_event(self, event):
        if event.src_path.lower().endswith(".docx"):
            print("📄 Modifica rilevata. Aggiorno...")
            if self.generator.load_all_files():
                self.generator.generate_csv_matrix()
                self.generator.generate_word_connections()
                print("✅ Aggiornato!")


def start_watcher(generator):
    event_handler = DocxChangeHandler(generator)
    observer = Observer()
    observer.schedule(event_handler, generator.watch_folder, recursive=False)
    observer.start()
    return observer


gen = CorrelationMatrixGenerator(WATCH_FOLDER)
# Cache per evitare timeout Gunicorn
global_df = None
global_connections = {}
last_update = 0



@app.route("/", methods=["GET", "POST"])
def home():
    global global_df, global_connections, last_update
    
    # Ricarica solo se >5min dall'ultima o primo avvio
    if time.time() - last_update > 300 or global_df is None:
        print("🔄 Ricaricamento dati...")
        gen.load_all_files()
        global_df, _ = gen.generate_csv_matrix()
        global_connections = gen.generate_word_connections()
        last_update = time.time()
        print("✅ Dati aggiornati")
    
    df = global_df
    gen.connections = global_connections  # per UI
    user_query = request.form.get("query", "").strip()
    chapter_query = request.form.get("chapter_query", "").strip()
    phrase_query = request.form.get("phrase_query", "").strip()


    if df is None:
        return "<h1>Nessun file DOCX trovato.</h1>"

    # -------------------------
    # Global index table
    # -------------------------
    global_rows = []
    for word, places in gen.connections.items():
        for doc_id, chap_title in places:
            global_rows.append({
                "PAROLA_CHIAVE": word,
                "POSIZIONE": f"{doc_id}: {chap_title}",
                "OCCORRENZE": len(places)
            })

    if global_rows:
        connections_df = pd.DataFrame(global_rows).sort_values(
            ["OCCORRENZE", "PAROLA_CHIAVE"],
            ascending=[False, True]
        ).reset_index(drop=True)

        connections_html = connections_df.to_html(
            index=False,
            table_id="connectionsTable",
            escape=False,
            classes="crossref-table"
        )
    else:
        connections_html = (
            '<p style="text-align:center;color:#666;padding:40px;">'
            '🔍 Nessuna parola trovata. Verifica i file .docx!</p>'
        )

    # Matrix HTML
    matrix_html = df.to_html(
        table_id="matrixTable",
        classes="matrix-table",
        escape=False,
        float_format=lambda x: f"{x:.3f}"
    )

    # -------------------------
    # Semantic query section
    # -------------------------
    query_results = gen.search_query_connections(user_query) if user_query else []
    if query_results:
        blocks = []
        for res in query_results:
            corr_rows = "".join(
                f"<tr><td>{c['DOCUMENTO']}</td><td>{c['SIMILARITÀ']:.3f}</td></tr>"
                for c in res["CORRELAZIONI"]
            )
            blocks.append(f"""
                <div class="query-block">
                    <h3>📍 Capitolo: <em>{res['PAESE_CAPITOLO']}</em></h3>
                    <p><strong>Parole query:</strong> {res['PAROLE_QUERY']}</p>
                    <table class="query-table">
                        <thead><tr><th>Articolo correlato</th><th>Similarità TF‑IDF</th></tr></thead>
                        <tbody>{corr_rows}</tbody>
                    </table>
                </div>
            """)
        query_html = "".join(blocks)
    elif user_query:
        query_html = "<p>❌ Nessun capitolo trovato che contenga tutte le parole della query.</p>"
    else:
        query_html = "<p>🔍 Inserisci 2+ parole per vedere correlazioni tra articoli.</p>"

    # -------------------------
    # Chapter-specific phrase output (always in your format)
    # -------------------------
    if not chapter_query:
        chapter_html = "<p>📚 Inserisci un capitolo. (Opzionale) Inserisci 1+ frasi/composti separati da ';' oppure a capo.</p>"
    else:
        reports = gen.analyze_chapter_specific_phrase_logic(
            chapter_query=chapter_query,
            phrase_query=phrase_query,
            top_n_links=12,
            min_score=0.06,
            alpha_phrase=0.70
        )

        if not reports:
            chapter_html = "<p>❌ Capitolo non trovato.</p>"
        else:
            pre_blocks = []
            for rep in reports:
                lines = rep["LINES"]
                # formato richiesto:
                # LABEL | first_location,
                # | other_location,
                # | ...
                out_lines = []
                first_loc = lines[0]["LOCATION"]
                out_lines.append(f"{rep['LABEL']} | {first_loc},")

                for idx, item in enumerate(lines[1:], start=1):
                    loc = item["LOCATION"]
                    comma = "," if idx < len(lines) - 1 else ""
                    out_lines.append(f"| {loc}{comma}")

                pre_blocks.append(f"<pre class='mono prebox'>{chr(10).join(out_lines)}</pre>")

            chapter_html = "".join(pre_blocks)

    html_content = f"""
<!DOCTYPE html>
<html>
<head>
    <title>Swallow Cross-Reference Intelligence</title>
    <meta http-equiv="refresh" content="3600">
    <style>
        body {{
            font-family: 'Segoe UI', Arial, sans-serif;
            margin: 40px;
            line-height: 1.5;
            background: #f5f5f5;
        }}
        h1 {{ color: #1e3a8a; text-align: center; }}
        h2 {{
            color: #1e40af;
            border-bottom: 3px solid #dbeafe;
            padding-bottom: 10px;
            margin-top: 40px;
        }}

        table {{
            border-collapse: collapse;
            width: 100%;
            font-size: 13px;
            margin: 15px 0;
            box-shadow: 0 4px 6px rgba(0,0,0,0.1);
            border-radius: 8px;
            overflow: hidden;
            background: white;
        }}
        th {{
            padding: 12px 15px;
            font-weight: 600;
            text-align: left;
        }}
        td {{
            padding: 12px 15px;
            border-bottom: 1px solid #e5e7eb;
            vertical-align: top;
        }}

        .stats {{
            background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
            color: white;
            padding: 20px;
            border-radius: 12px;
            margin: 20px 0;
            text-align: center;
        }}

        input[type="text"] {{
            width: 100%;
            padding: 14px;
            margin: 10px 0;
            border: 2px solid #d1d5db;
            border-radius: 8px;
            font-size: 15px;
            box-sizing: border-box;
            background: #fff;
            transition: 0.2s;
        }}
        input[type="text"]:focus {{
            border-color: #3b82f6;
            box-shadow: 0 0 0 3px rgba(59,130,246,0.1);
            outline: none;
        }}
        button[type="submit"] {{
            padding: 14px 24px;
            background: #4f46e5;
            color: white;
            border: none;
            border-radius: 8px;
            font-weight: 600;
            margin: 10px 0;
            box-shadow: 0 4px 6px rgba(79,70,229,0.3);
            transition: 0.2s;
            cursor: pointer;
        }}
        button[type="submit"]:hover {{
            background: #3730a3;
            transform: translateY(-1px);
        }}
        a {{
            display: inline-block;
            padding: 14px 24px;
            background: #3b82f6;
            color: white;
            text-decoration: none;
            border-radius: 8px;
            font-weight: 600;
            margin: 15px 0;
            box-shadow: 0 4px 6px rgba(59,130,246,0.3);
            transition: 0.2s;
        }}
        a:hover {{ background: #2563eb; transform: translateY(-1px); }}

        .matrix-table {{ font-size: 11px; }}
        .matrix-table th {{ background: #334155 !important; color: white; }}
        .matrix-table td {{ color: #475569; text-align: center; }}

        .crossref-table th:nth-child(1) {{
            background: linear-gradient(135deg, #fbbf24, #f59e0b);
            color: #1f2937;
            width: 25%;
        }}
        .crossref-table th:nth-child(2) {{
            background: linear-gradient(135deg, #10b981, #059669);
            color: white;
            width: 65%;
        }}
        .crossref-table th:nth-child(3) {{
            background: linear-gradient(135deg, #8b5cf6, #7c3aed);
            color: white;
            width: 10%;
            text-align: center;
        }}
        .crossref-table tr:nth-child(even) {{ background: #f8fafc; }}
        .crossref-table tr:hover {{ background: #e0f2fe !important; }}

        .connections {{
            max-height: 650px;
            overflow: auto;
            border-radius: 12px;
        }}
        .matrix-container {{
            max-height: 600px;
            overflow: auto;
            margin-top: 15px;
        }}

        .query-section {{
            background: linear-gradient(135deg, #eef2ff, #e0e7ff);
            padding: 25px;
            border-radius: 12px;
            margin: 20px 0;
            border: 2px solid #c7d2fe;
        }}
        .query-table th {{
            background: linear-gradient(135deg, #4f46e5, #3730a3);
            color: white;
        }}
        .query-block {{
            margin-top: 20px;
            padding: 20px;
            background: white;
            border-radius: 12px;
            box-shadow: 0 2px 8px rgba(0,0,0,0.08);
        }}

        .mono {{
            font-family: ui-monospace, SFMono-Regular, Menlo, Monaco, Consolas, "Liberation Mono", "Courier New", monospace;
            font-size: 12.5px;
        }}
        .prebox {{
            background: #0b1220;
            color: #e5e7eb;
            padding: 16px;
            border-radius: 12px;
            overflow: auto;
            white-space: pre;
            line-height: 1.4;
            border: 1px solid rgba(255,255,255,0.08);
            margin-top: 12px;
        }}
    </style>

    <script>
        function filterConnectionsTable() {{
            let input = document.getElementById('searchBox').value.toUpperCase();
            let table = document.getElementById('connectionsTable');
            if (!table) return;
            let trs = table.getElementsByTagName('tr');
            for (let i = 1; i < trs.length; i++) {{
                let txt = trs[i].textContent || trs[i].innerText;
                trs[i].style.display = txt.toUpperCase().indexOf(input) > -1 ? '' : 'none';
            }}
        }}

        function filterMatrix() {{
            let input = document.getElementById('matrixFilter').value.toUpperCase();
            let table = document.getElementById('matrixTable');
            if (!table) return;
            let trs = table.getElementsByTagName('tr');

            for (let i = 1; i < trs.length; i++) {{
                let row = trs[i];
                let rowHeader = row.cells[0].innerText.toUpperCase();
                let showRow = rowHeader.indexOf(input) > -1 || input === '';
                row.style.display = showRow ? '' : 'none';
            }}

            let headerRow = trs[0];
            for (let j = 1; j < headerRow.cells.length; j++) {{
                let colHeader = headerRow.cells[j].innerText.toUpperCase();
                let showCol = colHeader.indexOf(input) > -1 || input === '';
                headerRow.cells[j].style.display = showCol ? '' : 'none';
                for (let i = 1; i < trs.length; i++) {{
                    if (trs[i].cells[j]) {{
                        trs[i].cells[j].style.display = showCol ? '' : 'none';
                    }}
                }}
            }}
        }}
    </script>
</head>

<body>
    <h1>🧠 Swallow Cross-Reference Intelligence</h1>

    <div class="stats">
        <strong>📊 LIVE:</strong> {len(df)} documenti |
        {len(gen.connections)} parole indicizzate |
        {len(global_rows)} occorrenze |
        <em>Aggiornato: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</em>
    </div>

    <div class="query-section">
        <form method="POST">
            <label for="query"><strong>🔍 Cerca correlazioni tra articoli (Multi-parola)</strong><br>
            <small>Qui la query viene spezzata in parole (modalità “classica”).</small></label>
            <input type="text" id="query" name="query" value="{user_query}" placeholder="Es: volcano microclimate tourism">
            <button type="submit">Analizza query semantica</button>
        </form>
        <div style="margin-top:15px;">{query_html}</div>
    </div>

    <div class="query-section" style="background: linear-gradient(135deg, #f0fdf4, #dcfce7); border-color: #86efac;">
        <form method="POST">
            <label for="chapter_query"><strong>📚 Analizza capitolo specifico (frasi/composti non scorporati)</strong><br>
            <small>Il filtro sotto è una o più frasi/composti: separa con “;” oppure a capo.</small></label>

            <input type="text" id="chapter_query" name="chapter_query"
                   value="{chapter_query}"
                   placeholder="Capitolo (es: Stromboli / Diwali / Mount Fuji)">

            <input type="text" id="phrase_query" name="phrase_query"
                   value="{phrase_query}"
                   placeholder="Frase/composto (es: mount fuji; plate tectonics; volcanic landscape)">

            <button type="submit" style="background:#059669;">📌 Mostra output</button>
        </form>

        <div style="margin-top:15px;">{chapter_html}</div>
    </div>

    <h2>📈 Matrice Correlazione TF-IDF (documenti)</h2>
    <input type="text" id="matrixFilter" onkeyup="filterMatrix()" placeholder="🔍 Filtra documenti nella matrice...">
    <a href="/download_csv">⬇️ Scarica CSV Matrice</a>
    <div class="matrix-container">{matrix_html}</div>

    <h2>🔗 Indice parole (globale)</h2>
    <input type="text" id="searchBox" onkeyup="filterConnectionsTable()" placeholder="🔍 Filtra tabella parole/posizioni...">
    <div class="connections">{connections_html}</div>
</body>
</html>
"""
    return html_content


@app.route("/download_csv")
def download_csv():
    if os.path.exists(gen.csv_path):
        return send_file(
            gen.csv_path,
            mimetype="text/csv",
            as_attachment=True,
            download_name="swallow_correlation_matrix.csv"
        )
    return "<h1>❌ File CSV non trovato. Genera prima la matrice.</h1>"


if __name__ == "__main__":
    debug = os.environ.get("FLASK_DEBUG") == "1"
    port = int(os.environ.get("CONTAINER_PORT", "5000"))
    app.run(debug=debug, host="0.0.0.0", port=port, use_reloader=False)
